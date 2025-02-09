import streamlit as st
import docx
import tempfile
import os
from datetime import date
from docx.shared import Pt
import datetime
from docx import Document
# Function to replace placeholders in a Word document
TEMPLATE_PATH = os.path.abspath("testdoc.docx")
from docx import Document
from docx.shared import Pt

def replace_placeholders(doc_path, replacements):
    doc = Document(doc_path)

    # Function to replace placeholders in a paragraph, even if split across runs
    def replace_in_paragraph(para, replacements):
        if not para.runs:
            return  # Skip empty paragraphs

        # Store the font size of the first run
        first_run = para.runs[0]
        font_size = first_run.font.size or Pt(8)  # Default to 12pt if None
        is_bold = first_run.bold
        # Combine all runs into a single string
        full_text = "".join(run.text for run in para.runs)

        # Replace placeholders in the full text
        for key, value in replacements.items():
            full_text = full_text.replace(key, str(value))

        # Clear the paragraph and rebuild it with the replaced text
        para.clear()
        new_run = para.add_run(full_text)

        # Set only the font size
        new_run.font.size = font_size
        new_run.font.bold = is_bold

    # Replace placeholders in all paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    # Replace placeholders in tables, headers, and footers
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)

    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            for para in header.paragraphs:
                replace_in_paragraph(para, replacements)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            for para in footer.paragraphs:
                replace_in_paragraph(para, replacements)

    return doc
# def replace_placeholders(doc_path, replacements):
    doc = Document(doc_path)

    def replace_in_paragraph(para, replacements):
        for run in para.runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, str(value))
                    run.font.bold = run.bold  # Retain bold if it was originally bold
                    run.font.size = run.font.size or Pt(8)  # Preserve font size

    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    # Replace placeholders in tables, headers, and footers
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)

    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            for para in header.paragraphs:
                replace_in_paragraph(para, replacements)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            for para in footer.paragraphs:
                replace_in_paragraph(para, replacements)

    return doc

# Streamlit UI
st.title("📄 Modify Word Template - Income Tax Assessment")


totalAtoF=0

st.write("🔹 Enter Required Information")
    
with st.expander("Institute & Assessment Details"):
    institute_name = st.text_input("Institute Name")
    fin_year = st.text_input("Financial Year")
    assess_year = st.text_input("Assessment Year")
    pan = st.text_input("PAN No.")
    aadhar = st.text_input("Aadhar No.")
    unique_id = st.text_input("Unique ID")
    treasury_code = st.text_input("Treasury Code")
    
    # Employee Details
with st.expander("Employee Details"):
        name = st.text_input("Employee Name")
        father_name = st.text_input("Father's Name")
        min_date = datetime.date(1920, 1, 1)
        max_date = datetime.date.today()
        dob = st.date_input("Date of Birth", min_value=min_date, max_value=max_date)
        designation = st.text_input("Designation")
        address = st.text_area("Residence Address")
    
with st.expander("Income Details "):
        basic_pay = st.number_input("Basic/Spl. Pay", min_value=0.0, key="basic_pay")
        agp = st.number_input("A.G.P.", min_value=0.0, key="agp",value=0.0)
        da = st.number_input("Dearness Allowance", min_value=0.0, key="da")
        hra = st.number_input("H.R.A.", min_value=0.0, key="hra_1",value=0.0)
        arrears = st.number_input("Other (Arrears)", min_value=0.0, key="arrears",value=0.0)
        others = st.number_input("Other Allowances", min_value=0.0, key="hra_2",value=0.0)  # Rename HRA duplicate
        incomeOtherSources = st.number_input("Income From other Sources", min_value=0.0, key="income_other_sources",value=0.0)
        houseRent = st.number_input("Actual House Rent Received", min_value=0.0, key="house_rent",value=0.0)
        excessRentPaid = st.number_input("Rent paid in excess of 10% of salary", min_value=0.0, key="excess_rent_paid",value=0.0)
        houseLoanInterest = st.number_input("Less Interest of Loan upto Rs. 2,00,000/-", min_value=0.0, key="house_loan_interest",value=0.0)
        standard_deduction = st.number_input("Standard Deduction (Upto 50000)", min_value=0.0, key="standard_deduction")
        professional_tax = st.number_input("Professional Tax", min_value=0.0, key="professional_tax",value=0.0)
        totalAtoF=basic_pay + agp + da + hra + arrears +others
        total_income = basic_pay + agp + da + hra + arrears + incomeOtherSources - houseLoanInterest-(min(houseRent,excessRentPaid))-standard_deduction - professional_tax
        st.write(f"Total Income: {total_income}")
    # Deduction Section
with st.expander("80 C Deductions"):
        lic = st.number_input("L.I.C.", min_value=0.0,value=0.0)
        gpf = st.number_input("G.P.F./P.P.F./G.I.S.", min_value=0.0,value=0.0)
        nsc = st.number_input("N.S.C.", min_value=0.0,value=0.0)
        nsc_int = st.number_input("N.S.C. Accrued Interest", min_value=0.0,value=0.0)
        home_loan_p = st.number_input("Repayment of House Loan (Principal)", min_value=0.0,value=0.0)
        tuition_fee = st.number_input("Tuition Fees (Up to 2 Children)", min_value=0.0,value=0.0)
        bank_fdr = st.number_input("Bank FDR (5-year)", min_value=0.0,value=0.0)
        mutual_fund = st.number_input("Mutual Fund (3-year)", min_value=0.0,value=0.0)
        nps = st.number_input("New Pension Scheme", min_value=0.0,value=0.0)
        specify_other = st.text_input("Other Deduction (Specify)")
        other = st.number_input("Other Deduction Amount", min_value=0.0,value=0.0)
        total_deductions = lic + gpf + nsc + nsc_int + home_loan_p + tuition_fee + bank_fdr + mutual_fund + nps + other 
        total_deductions=min(150000,total_deductions)
        st.write(f"Deductions under 80C: {total_deductions}")

with st.expander("Other Deductions"):
        ccd80 = st.number_input("80CCD Investment in New Pension Scheme (other than 80C upto 50000/-", min_value=0.0,value=0.0)
        d80 = st.number_input("80D Medical Insurance premium upto 25000/- (for Sr.Citizen) 50000/-", min_value=0.0,value=0.0)
        dd80 = st.number_input("80DD Dependent Handicapped exemption upto 75000/-", min_value=0.0,value=0.0)
        e80 = st.number_input("80E Interest of education loan", min_value=0.0,value=0.0)
        u80 = st.number_input("80U Physically Handicapped Exemption upto 75000/-", min_value=0.0,value=0.0)
        ttb80 = st.number_input("80TTB Interest of Saving & FD for Sr. Citizen upto 50000/-", min_value=0.0,value=0.0)
        otherDeductions=ccd80+d80+dd80+e80+u80+ttb80
        st.write(f"Other Deductions: {otherDeductions}")
    # Tax Calculation
overallDeduction=total_deductions+otherDeductions
st.write(f"Total Deductions: {overallDeduction}")
net_income = total_income - overallDeduction
st.write(f"Net Taxable Income: {net_income}")

per5=0
per20=0
per30=0
tax=0
rebate=0
totalTax=0
educationCess=0
relief89 =0
per5s=0
today = date.today()
age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
if(net_income<=250000):
    tax=0
elif(net_income>250000 and net_income<=500000 and age<=65 ):
    per5=round((net_income-250000)*0.05)
    tax=per5
    rebate=12500
elif(net_income>250000 and net_income<=500000 and age>65  ):
    per5s=round((net_income-300000)*0.05)
    per5=0
    tax=per5
    rebate =12500
elif(net_income>500000 and net_income<=1000000 ):
    per5=12500
    if(age>65):per5s=12500
    per20=round((net_income-500000)*0.2)
    tax=per5+per20
elif(net_income>1000000 ):
    per5=12500
    if(age>65):per5s=12500
    per20=100000
    per30=round((net_income-1000000)*0.3)
    tax=per5+per20+per30
    
balance=tax-rebate
payableTax=balance;
if(payableTax<0):
    payableTax=0
educationCess=round(0.04*payableTax)
totalTax=tax+educationCess
   
   

    # Advance Tax Paid
with st.expander("Advance Tax Paid"):
    months = ["apr", "may", "jun", "jul", "aug", "sep", 
                "oct", "nov", "dec", "jan", "feb", "mar"]
        
        # Dictionary to store tax paid for each month
    tax_paid = {}

    for month in months:
        tax_paid[month] = st.number_input(f"{month} Tax Paid", min_value=0.0,value=0.0)

        # Calculate total tax paid
    total_tax_paid = sum(tax_paid.values())
        
        # Display total tax paid
    st.write(f"**Total Tax Paid:** {total_tax_paid}")

    # Add tax values to the replacements dictionary for Word document
replacements = {
        "{{totalTaxPaid}}": str(total_tax_paid),  # Store total tax paid
    }
toPay=totalTax-total_tax_paid
with st.expander("Tax Calculation"):
        st.write(f"Total Tax On Income: {tax}")   
        st.write(f"Tax Rebate of Rs. 12500/- (u/s 87A for NTI ) Taxable Income is Rs.500000/ or Less: {rebate}")  
        st.write(f"Balance Tax payable: {payableTax}")  
        st.write(f"Education Cess (4%): {educationCess}")  
        st.write(f"Total Tax Payable: {totalTax}")  
        st.write(f"Total Tax paid: {total_tax_paid}")  
    
    # Declaration
with st.expander("Declaration"):
        place = st.text_input("Place")
        date = st.date_input("Date")
        st.write(f"I, {name}, working as {designation}, declare that the information provided is correct to the best of my knowledge.")

     
    # Button to generate modified document
if st.button("Generate Document"):
    if name and designation:
        # Save uploaded template to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_template:
            with open(TEMPLATE_PATH, "rb") as template_file:
                temp_template.write(template_file.read())  # Write actual file content
            temp_template_path = temp_template.name

        # Define replacements correctly
        replacements = {
            "{{instituteName}}": institute_name,
            "{{finYear}}": fin_year,
            "{{assessYear}}": assess_year,
            "{{pan}}": pan,
            "{{aadhar}}": aadhar,
            "{{uniqueId}}": unique_id,
            "{{treasuryCode}}": treasury_code,
            "{{name}}": name,
            "{{fatherName}}": father_name,
            "{{dob}}": str(dob),
            "{{designation}}": designation,
            "{{address}}": address,
            "{{basicPay}}": str(round(basic_pay)),
            "{{agp}}": str(round(agp)),
            "{{da}}": str(round(da)),
            "{{hra}}": str(round(hra)),
            "{{netIncome}}": str(round(net_income)),
            "{{totalTaxPaid}}": str(round(total_tax_paid)),
            "{{place}}": place,
            "{{date}}": str(date),
            "{{incomeOtherSources}}": str(round(incomeOtherSources)),
            "{{houseLoanInterest}}": str(round(houseLoanInterest)),
            "{{5per}}": str(round(per5)),
            "{{5perS}}": str(round(per5s)),
            "{{20per}}": str(round(per20)),
            "{{30per}}": str(round(per30)),
            "{{totalTax}}": str(round(tax)),
            "{{rebate}}": str(round(rebate)),
            "{{balance}}": str(round(balance)),
            "{{educationCess}}": str(round(educationCess)),
            "{{totalTaxPayable}}": str(round(totalTax)),
            "{{relief89}}": str(round(relief89)),
            "{{taxPaid}}": str(round(total_tax_paid)),
            "{{houseRent}}": str(round(houseRent)),
            "{{excessRentPaid}}": str(round(excessRentPaid)),
            "{{standard}}": str(round(standard_deduction)),
            "{{proffTax}}": str(round(professional_tax)),
            "{{grossTotal}}": str(round(total_income)),
            "{{lic}}": str(round(lic)),
            "{{gpf}}": str(round(gpf)),
            "{{nsc}}": str(round(nsc)),
            "{{nscInt}}": str(round(nsc_int)),
            "{{homeloanP}}": str(round(home_loan_p)),
            "{{tuitionfee}}": str(round(tuition_fee)),
            "{{bankfdr}}": str(round(bank_fdr)),
            "{{mutualFund}}": str(round(mutual_fund)),
            "{{nps}}": str(round(nps)),
            "{{other}}": str(round(other)),
            "{{specifyOther}}": str(specify_other),
            "{{80ccd}}": str(round(ccd80)),
            "{{80d}}": str(round(d80)),
            "{{80dd}}": str(round(dd80)),
            "{{80e}}": str(round(e80)),
            "{{80u}}": str(round(u80)),
            "{{80ttb}}": str(round(ttb80)),
            "{{deduction}}": str(round(overallDeduction)),
            "{{others}}": str(round(others)),
            "{{arrear}}": str(round(arrears)),
            "{{totalAtoF}}": str(round(totalAtoF)),
            "{{topay}}": str(round(toPay)),
            "{{80c}}": str(round(total_deductions)),
            "{{mar}}":str(tax_paid["mar"]),
            "{{jan}}": str(tax_paid["jan"]),
            "{{feb}}": str(tax_paid["feb"]),
            "{{mar}}": str(tax_paid["mar"]),
            "{{apr}}": str(tax_paid["apr"]),
            "{{may}}": str(tax_paid["may"]),
            "{{jun}}": str(tax_paid["jun"]),
            "{{jul}}": str(tax_paid["jul"]),
            "{{aug}}": str(tax_paid["aug"]),
            "{{sep}}": str(tax_paid["sep"]),
            "{{oct}}": str(tax_paid["oct"]),
            "{{nov}}": str(tax_paid["nov"]),
            "{{dec}}": str(tax_paid["dec"]),
        }

        # Modify the Word file
        modified_doc = replace_placeholders(temp_template_path, replacements)

        # Save the modified document
        output_path = "output.docx"
        modified_doc.save(output_path)

        # Provide download link
        with open(output_path, "rb") as f:
            st.download_button("📥 Download Modified Word File", f, file_name="Modified_Document.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Clean up temporary files
        os.remove(temp_template_path)
        os.remove(output_path)

    else:
        st.error("⚠️ Please fill in required fields.")